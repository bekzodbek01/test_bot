import random
import os
import time
from aiogram import Bot, Dispatcher, F
from aiogram.enums import PollType
from aiogram.filters import Command
from aiogram.types import (
    Message,
    KeyboardButton,
    PollAnswer
)

import json
from aiogram.utils.keyboard import ReplyKeyboardBuilder
from aiogram.fsm.storage.memory import MemoryStorage
import sys
from docx import Document
import asyncio

if sys.platform.startswith("win"):
    asyncio.set_event_loop_policy(
        asyncio.WindowsSelectorEventLoopPolicy()
    )
import logging

from dotenv import load_dotenv

load_dotenv()
TOKEN = os.getenv("BOT_TOKEN")
logging.basicConfig(level=logging.INFO)

DATA_FOLDER = "data"

QUESTIONS_PER_TEST = 20

bot = Bot(TOKEN)

dp = Dispatcher(storage=MemoryStorage())

users_test = {}
user_results = {}
leaderboard_data = {}

LEADERBOARD_FILE = "leaderboard.json"


# ==================================================
# CLEAN TEXT
# ==================================================

def clean_text(text):
    return " ".join(str(text).split()).strip()


# ==================================================
# SAVE / LOAD LEADERBOARD
# ==================================================

def load_leaderboard():
    global leaderboard_data

    if os.path.exists(LEADERBOARD_FILE):

        with open(
                LEADERBOARD_FILE,
                "r",
                encoding="utf-8"
        ) as f:

            leaderboard_data = json.load(f)

    else:
        leaderboard_data = {}


def save_leaderboard():
    with open(
            LEADERBOARD_FILE,
            "w",
            encoding="utf-8"
    ) as f:
        json.dump(
            leaderboard_data,
            f,
            ensure_ascii=False,
            indent=4
        )


# ==================================================
# MAIN MENU
# ==================================================

def main_menu():
    kb = ReplyKeyboardBuilder()

    kb.row(
        KeyboardButton(text="🚀 Testni boshlash")
    )

    kb.row(
        KeyboardButton(text="📊 Natijalar"),
        KeyboardButton(text="🏆 Reyting")
    )

    kb.row(
        KeyboardButton(text="📚 Fanlar")
    )

    return kb.as_markup(resize_keyboard=True)


# ==================================================
# LOAD QUESTIONS
# ==================================================


# =========================================
# LOAD QUESTIONS
# =========================================

def load_questions():
    questions = []

    for filename in os.listdir(DATA_FOLDER):

        if filename.startswith("~$"):
            continue

        if not filename.endswith(".docx"):
            continue

        path = os.path.join(DATA_FOLDER, filename)

        print(f"Yuklanmoqda: {filename}")

        try:

            doc = Document(path)

            # =====================================
            # TABLE O'QISH
            # =====================================

            for table in doc.tables:

                for row in table.rows:

                    try:

                        cells = row.cells

                        # normal row
                        if len(cells) < 5:
                            continue

                        texts = []

                        for cell in cells:

                            txt = " ".join(
                                clean_text(p.text)
                                for p in cell.paragraphs
                                if clean_text(p.text)
                            )

                            txt = clean_text(txt)

                            if txt:
                                texts.append(txt)

                        # minimum 5 ta ustun
                        if len(texts) < 5:
                            continue

                        # =====================================
                        # SAVOL
                        # =====================================

                        question = clean_text(texts[0])

                        if not question:
                            continue

                        # HEADER FILTER
                        # HEADERLARNI FILTER QILISH
                        bad_headers = [
                            "test savoli",
                            "to‘g‘ri javob",
                            "muqobil javob",
                            "universitet",
                            "business and science",
                            "o‘quv yili",
                            "ta’lim yo‘nalishi",
                            "yakuniy davlat attestatsiyasi",
                        ]

                        q_lower = question.lower()

                        # faqat headerlarni skip qil
                        if any(x in q_lower for x in bad_headers):
                            continue

                        # juda qisqa bo‘lsa skip
                        if len(question.strip()) < 3:
                            continue

                        # =====================================
                        # JAVOBLAR
                        # =====================================

                        answers = []

                        for ans in texts[1:5]:

                            ans = clean_text(ans)

                            if not ans:
                                ans = "Variant mavjud emas"

                            # Telegram limit
                            ans = ans[:95]

                            answers.append(ans)

                        # duplicate variantlarni olib tashlash
                        answers = list(dict.fromkeys(answers))

                        # minimum 2 ta javob
                        if len(answers) < 2:
                            continue

                        # 4 tagacha to'ldirish
                        while len(answers) < 4:
                            answers.append("Variant mavjud emas")

                        answers = answers[:4]

                        # =====================================
                        # SAVE
                        # =====================================

                        questions.append({
                            "subject": filename.replace(".docx", ""),
                            "question": question[:300],
                            "correct": answers[0],
                            "answers": answers
                        })

                    except Exception as e:
                        print("TABLE ERROR:", e)

        except Exception as e:

            print("XATO:", filename)
            print(e)

    print(f"JAMI SAVOLLAR: {len(questions)}")

    return questions


ALL_QUESTIONS = load_questions()
load_leaderboard()


# ==================================================
# START
# ==================================================

@dp.message(Command("start"))
async def start(message: Message):
    text = f"""
🎓 PROFESSIONAL QUIZ BOT

📚 Jami savollar:
{len(ALL_QUESTIONS)}

✅ Quiz Poll
✅ 20 random savol
✅ 25 sekund timer
✅ Reyting
✅ Natijalar
"""

    await message.answer(
        text,
        reply_markup=main_menu()
    )


# ==================================================
# TEST MENU
# ==================================================

@dp.message(F.text == "🚀 Testni boshlash")
async def test_menu(message: Message):
    counts = {}

    for q in ALL_QUESTIONS:
        subject = q["subject"]

        counts[subject] = counts.get(subject, 0) + 1

    text = "📚 Qaysi fandan test ishlaysiz?\n\n"

    kb = ReplyKeyboardBuilder()

    total = 0

    for subject, count in counts.items():
        total += count

        text += f"📘 {subject} — {count} ta savol\n"

        kb.row(
            KeyboardButton(
                text=f"📘 {subject}"
            )
        )

    text += f"\n🌍 Barcha fanlar — {total} ta savol"

    kb.row(
        KeyboardButton(text="🌍 Barcha fanlar")
    )

    kb.row(
        KeyboardButton(text="🛑 Testni tugatish")
    )

    await message.answer(
        text,
        reply_markup=kb.as_markup(
            resize_keyboard=True
        )
    )


# ==================================================
# SUBJECT TEST
# ==================================================

@dp.message(F.text.startswith("📘 "))
async def subject_test(message: Message):
    user_id = message.from_user.id

    subject_name = (
        message.text
        .replace("📘 ", "")
        .split(" — ")[0]
    )

    questions = [
        q for q in ALL_QUESTIONS
        if q["subject"] == subject_name
    ]

    if not questions:
        await message.answer(
            "❌ Savollar topilmadi"
        )
        return

    random.shuffle(questions)

    questions = questions[:QUESTIONS_PER_TEST]

    users_test[user_id] = {
        "questions": questions,
        "index": 0,
        "correct": 0,
        "wrong": 0,
        "skipped": 0,
        "chat_id": message.chat.id,
        "start_time": time.time(),
        "answered": False,
        "subjects": {},
        "full_name": message.from_user.full_name

    }

    await send_question(user_id)


# ==================================================
# ALL SUBJECTS TEST
# ==================================================

@dp.message(F.text == "🌍 Barcha fanlar")
async def all_subjects_test(message: Message):
    user_id = message.from_user.id

    questions = ALL_QUESTIONS.copy()

    random.shuffle(questions)

    questions = questions[:QUESTIONS_PER_TEST]

    users_test[user_id] = {
        "questions": questions,
        "index": 0,
        "correct": 0,
        "wrong": 0,
        "skipped": 0,
        "chat_id": message.chat.id,
        "start_time": time.time(),
        "answered": False,
        "subjects": {},
        "full_name": message.from_user.full_name
    }

    await send_question(user_id)


async def send_question(user_id):
    if user_id not in users_test:
        return

    data = users_test[user_id]

    if data["index"] >= len(data["questions"]):

        total = (
                data["correct"] +
                data["wrong"] +
                data["skipped"]
        )

        correct = data["correct"]

        percent = round(
            (correct / total) * 100,
            1
        ) if total > 0 else 0

        spent = int(
            time.time() - data["start_time"]
        )

        # FANLAR BO'YICHA
        # =====================================

        subjects_text = "\n📚 Fanlar bo'yicha:\n\n"

        for sub, stat in data["subjects"].items():
            subjects_text += (
                f"📘 {sub}\n"
                f"✅ {stat['correct']} | "
                f"❌ {stat['wrong']} | "
                f"⏭ {stat['skipped']}\n\n"
            )

        # =====================================
        # RESULT TEXT

        result_text = f"""
        😔 TEST NATIJALARI

        👤 User ID: {user_id}

        📊 Jami savol: {total}

        ✅ To'g'ri: {data['correct']}
        ❌ Noto'g'ri: {data['wrong']}
        ⏭ O'tkazilgan: {data['skipped']}

        🎯 Ball: {percent}%

        ⏱ Vaqt:
        {spent // 60} daq {spent % 60} son

        {subjects_text}
        """

        # =====================================
        # RESULT SAVE
        # =====================================

        user_results[user_id] = result_text

        # =====================================
        # GLOBAL REYTING
        # =====================================

        if str(user_id) not in leaderboard_data:
            leaderboard_data[str(user_id)] = {

                "name": data["full_name"],
                "correct": 0,
                "total": 0,
            }

        leaderboard_data[str(user_id)]["correct"] += data["correct"]

        leaderboard_data[str(user_id)]["total"] += total

        save_leaderboard()

        await bot.send_message(
            data["chat_id"],
            result_text,
            reply_markup=main_menu()
        )

        del users_test[user_id]

        return

    q = data["questions"][data["index"]]

    answers = []

    for a in q["answers"]:

        a = clean_text(a)

        if len(a) > 90:
            a = a[:90]

        answers.append(a)

    answers = list(dict.fromkeys(answers))

    if len(answers) < 2:
        data["index"] += 1

        await send_question(user_id)

        return

    correct_answer = clean_text(q["correct"])

    if len(correct_answer) > 90:
        correct_answer = correct_answer[:90]

    if correct_answer not in answers:
        answers[0] = correct_answer

    random.shuffle(answers)

    correct_index = answers.index(correct_answer)

    data["correct_index"] = correct_index
    data["answered"] = False

    text = (
        f"📝 {data['index'] + 1}/{len(data['questions'])} | "
        f"🚀 {q['subject']}\n\n"
        f"{q['question'][:250]}"
    )

    await bot.send_poll(
        chat_id=data["chat_id"],
        question=text[:300],
        options=answers[:4],
        type=PollType.QUIZ,
        correct_option_id=correct_index,
        is_anonymous=False,
        # open_period=25,# timer ochirish
        explanation=f"✅ To'g'ri: {correct_answer}"
    )


# ==================================================
# POLL ANSWER
# ==================================================

@dp.poll_answer()
async def poll_answer_handler(
        poll_answer: PollAnswer
):
    user_id = poll_answer.user.id

    if user_id not in users_test:
        return

    data = users_test[user_id]

    if data["answered"]:
        return

    data["answered"] = True

    # =====================================
    # CURRENT QUESTION
    # =====================================

    q = data["questions"][data["index"]]

    subject = q["subject"]

    # subject stat yaratish
    if subject not in data["subjects"]:
        data["subjects"][subject] = {
            "correct": 0,
            "wrong": 0,
            "skipped": 0
        }

    # =====================================
    # SKIPPED
    # =====================================

    if not poll_answer.option_ids:

        data["skipped"] += 1

        data["subjects"][subject]["skipped"] += 1

    else:

        selected = poll_answer.option_ids[0]

        # =================================
        # TO'G'RI
        # =================================

        if selected == data["correct_index"]:

            data["correct"] += 1

            data["subjects"][subject]["correct"] += 1

        # =================================
        # NOTO'G'RI
        # =================================

        else:

            data["wrong"] += 1

            data["subjects"][subject]["wrong"] += 1

    # =====================================
    # NEXT QUESTION
    # =====================================

    data["index"] += 1

    await asyncio.sleep(2)

    await send_question(user_id)


# ==================================================
# RESULTS
# ==================================================

@dp.message(F.text == "📊 Natijalar")
async def results(message: Message):
    user_id = message.from_user.id

    if user_id not in user_results:
        await message.answer(
            "❌ Siz hali test ishlamagansiz"
        )
        return

    await message.answer(
        user_results[user_id]
    )


# ==================================================
# GLOBAL REYTING
# ==================================================

@dp.message(F.text == "🏆 Reyting")
async def leaderboard(message: Message):
    # reyting bo'sh
    if not leaderboard_data:
        await message.answer(
            "❌ Reyting hali bo'sh"
        )
        return

    board = []

    # =====================================
    # DATA
    # =====================================

    for user_id, data in leaderboard_data.items():
        correct = data["correct"]
        total = data["total"]

        percent = round(
            (correct / total) * 100,
            1
        ) if total > 0 else 0

        board.append({
            "name": data["name"],
            "correct": correct,
            "total": total,
            "percent": percent
        })

    # =====================================
    # SORT
    # =====================================

    board.sort(
        key=lambda x: (
            x["correct"],
            x["percent"]
        ),
        reverse=True
    )

    # =====================================
    # TEXT
    # =====================================

    text = "🏆 GLOBAL REYTING\n\n"

    for i, user in enumerate(board[:20], start=1):
        text += (
            f"{i}. 👤 {user['name']}\n"
            f"✅ {user['correct']}/{user['total']} "
            f"({user['percent']}%)\n\n"
        )

    # =====================================
    # SEND
    # =====================================

    await message.answer(text)


# ==================================================
# FANLAR
# ==================================================

@dp.message(F.text == "📚 Fanlar")
async def subjects(message: Message):
    counts = {}

    total = 0

    for q in ALL_QUESTIONS:
        counts[q["subject"]] = (
                counts.get(q["subject"], 0) + 1
        )

        total += 1

    text = "📚 MAVJUD FANLAR\n\n"

    for subject, count in counts.items():
        text += f"📘 {subject} — {count} ta savol\n"

    text += f"\n📝 Jami: {total} ta test"

    await message.answer(text)


# ==================================================
# STOP TEST
# ==================================================


# ==================================================
# TESTNI TUGATISH
# ==================================================

@dp.message(F.text == "🛑 Testni tugatish")
async def stop_test(message: Message):
    user_id = message.from_user.id

    # test boshlanmagan
    if user_id not in users_test:
        await message.answer(
            "❌ Siz hali test boshlamagansiz",
            reply_markup=main_menu()
        )
        return

    data = users_test[user_id]

    total = (
            data["correct"] +
            data["wrong"] +
            data["skipped"]
    )

    # hech narsa ishlamagan
    if total == 0:
        await message.answer(
            "❌ Siz hali savol ishlamagansiz",
            reply_markup=main_menu()
        )
        return

    # =====================================
    # BALL
    # =====================================

    percent = round(
        (data["correct"] / total) * 100,
        1
    )

    spent = int(
        time.time() - data["start_time"]
    )

    # =====================================
    # FANLAR
    # =====================================

    subjects_text = "\n📚 Fanlar bo'yicha:\n\n"

    for sub, stat in data["subjects"].items():
        subjects_text += (
            f"📘 {sub}\n"
            f"✅ {stat['correct']} | "
            f"❌ {stat['wrong']} | "
            f"⏭ {stat['skipped']}\n\n"
        )

    # =====================================
    # RESULT
    # =====================================

    result_text = f"""
😔 TEST NATIJALARI

👤 User ID: {user_id}

📊 Jami savol: {total}

✅ To'g'ri: {data['correct']}
❌ Noto'g'ri: {data['wrong']}
⏭ O'tkazilgan: {data['skipped']}

🎯 Ball: {percent}%

⏱ Vaqt:
{spent // 60} daq {spent % 60} son

{subjects_text}
"""

    # saqlash
    user_results[user_id] = result_text

    # =====================================
    # GLOBAL REYTING
    # =====================================

    if str(user_id) not in leaderboard_data:
        leaderboard_data[str(user_id)] = {
            "name": message.from_user.full_name,
            "correct": 0,
            "total": 0,
        }

    leaderboard_data[str(user_id)]["correct"] += data["correct"]

    leaderboard_data[str(user_id)]["total"] += total

    save_leaderboard()

    await message.answer(
        result_text,
        reply_markup=main_menu()
    )

    # testni yopish
    del users_test[user_id]


# =========================================
# MAIN
# =========================================

async def main():
    while True:

        try:

            print("BOT ISHGA TUSHDI")
            print("SAVOLLAR:", len(ALL_QUESTIONS))

            await dp.start_polling(bot)

        except Exception as e:

            print("XATO:", e)
            print("5 sekunddan keyin qayta ulanadi...")

            await asyncio.sleep(5)


# =========================================
# START
# =========================================

if __name__ == "__main__":

    try:
        asyncio.run(main())

    except KeyboardInterrupt:
        print("Dastur yopildi")

