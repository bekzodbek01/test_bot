







import os
import sys
import time
import random
import logging
import asyncio
from dotenv import load_dotenv
from docx import Document

from aiogram import Bot, Dispatcher, F
from aiogram.enums import PollType
from aiogram.filters import Command
from aiogram.utils.keyboard import ReplyKeyboardBuilder
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.context import FSMContext
from aiogram.types import Message, ReplyKeyboardMarkup, KeyboardButton, PollAnswer

# Ma'lumotlar bazasi funksiyalarini import qilish
from utils.database import (
    clear_fake_users, is_admin_db, is_user, add_admin_db,
    save_result, del_admin_db, add_user_db, del_user_db,
    get_users, get_admins, get_leaderboard
)

# Startup DB tozalash
try:
    clear_fake_users()
except Exception as e:
    logging.warning(f"Bazada dastlabki tozalashda xatolik: {e}")

# Windows operatsion tizimi uchun Event Loop moslashuvi
if sys.platform.startswith("win"):
    asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())

# Logging va Environment sozlamalari
logging.basicConfig(level=logging.INFO)
load_dotenv()

TOKEN = os.getenv("BOT_TOKEN")
SUPER_ADMIN = int(os.getenv("SUPER_ADMIN", "0"))
DATA_FOLDER = "data"
QUESTIONS_PER_TEST = 20

bot = Bot(token=TOKEN)
dp = Dispatcher(storage=MemoryStorage())


# ==================================================
# FSM STATES (HOLATLAR)
# ==================================================
class AddAdmin(StatesGroup): wait = State()


class DelAdmin(StatesGroup): wait = State()


class AddUser(StatesGroup): wait = State()


class DelUser(StatesGroup): wait = State()


class SearchState(StatesGroup): wait_keyword = State()  # Mukammal qidiruv uchun holat


# Keshlash ob'ektlari
users_test = {}
user_results = {}


# ==================================================
# YORDAMCHI FUNKSIYALAR va RUXSATLAR
# ==================================================
def is_super(uid: int) -> bool:
    return uid == SUPER_ADMIN


def is_admin(uid: int) -> bool:
    return uid == SUPER_ADMIN or is_admin_db(uid)


def clean_text(text: str) -> str:
    return " ".join(str(text).split()).strip()


# ==================================================
# KLAVIATURALAR (image_da6fc2.png dizayniga moslangan)
# ==================================================
def main_menu(user_id: int) -> ReplyKeyboardMarkup:
    kb = [
        [KeyboardButton(text="🚀 Testni boshlash")],
        [KeyboardButton(text="📊 Natijalar"), KeyboardButton(text="🏆 Reyting")],
        [KeyboardButton(text="📚 Fanlar"), KeyboardButton(text="🔍 Qidiruv")]  # Rasmga mos yonma-yon dizayn
    ]
    if is_admin(user_id):
        kb.append([KeyboardButton(text="👑 Admin Panel")])
    return ReplyKeyboardMarkup(keyboard=kb, resize_keyboard=True)


def admin_menu() -> ReplyKeyboardMarkup:
    return ReplyKeyboardMarkup(keyboard=[
        [KeyboardButton(text="👥 Users"), KeyboardButton(text="🛡 Admin")],
        [KeyboardButton(text="🏠 Home")]
    ], resize_keyboard=True)


def users_menu() -> ReplyKeyboardMarkup:
    return ReplyKeyboardMarkup(keyboard=[
        [KeyboardButton(text="➕ User"), KeyboardButton(text="❌ Del User")],
        [KeyboardButton(text="📋 User List")],
        [KeyboardButton(text="⬅ Back")]
    ], resize_keyboard=True)


def admins_menu() -> ReplyKeyboardMarkup:
    return ReplyKeyboardMarkup(keyboard=[
        [KeyboardButton(text="➕ Admin"), KeyboardButton(text="❌ Del Admin")],
        [KeyboardButton(text="📋 Admin List")],
        [KeyboardButton(text="⬅ Back")]
    ], resize_keyboard=True)


# ==================================================
# DOCX STRUKTURASINI O'QISH (SAVOLLAR BAZASI)
# ==================================================
def load_questions():
    questions = []
    if not os.path.exists(DATA_FOLDER):
        os.makedirs(DATA_FOLDER)
        return questions

    for filename in os.listdir(DATA_FOLDER):
        if filename.startswith("~$") or not filename.endswith(".docx"):
            continue

        path = os.path.join(DATA_FOLDER, filename)
        try:
            doc = Document(path)
            for table in doc.tables:
                for row in table.rows:
                    try:
                        cells = row.cells
                        if len(cells) < 5:
                            continue

                        texts = [clean_text(" ".join(p.text for p in cell.paragraphs if clean_text(p.text))) for cell in
                                 cells]
                        texts = [t for t in texts if t]

                        if len(texts) < 5:
                            continue

                        question = texts[0]
                        bad_headers = ["test savoli", "to‘g‘ri javob", "muqobil javob", "universitet",
                                       "business and science", "o‘quv yili", "ta’lim yo‘nalishi",
                                       "yakuniy davlat attestatsiyasi"]
                        if any(x in question.lower() for x in bad_headers) or len(question.strip()) < 3:
                            continue

                        answers = [clean_text(ans)[:95] for ans in texts[1:5]]
                        answers = list(dict.fromkeys(answers))  # nusxalarni o'chirish

                        while len(answers) < 4:
                            answers.append("Variant mavjud emas")
                        answers = answers[:4]

                        questions.append({
                            "subject": filename.replace(".docx", ""),
                            "question": question[:300],
                            "correct": answers[0],
                            "answers": answers
                        })
                    except Exception as e:
                        logging.error(f"Ustun tahlilida xatolik: {e}")
        except Exception as e:
            logging.error(f"Faylni o'qishda xatolik {filename}: {e}")

    print(f"✅ JAMI {len(questions)} TA SAVOL SIZNING 7 TA FANINGIZDAN YUKLANDI")
    return questions


ALL_QUESTIONS = load_questions()


# ==================================================
# START BUYRUG'I
# ==================================================
@dp.message(Command("start"))
async def start(message: Message):
    uid = message.from_user.id
    if uid != SUPER_ADMIN and not is_admin_db(uid) and not is_user(uid):
        await message.answer("⛔ Sizga ruxsat yo‘q\n\nBotdan foydalanish uchun:\n📩 @dasturchi_0101 ga murojaat qiling.")
        return

    if is_super(uid) or is_admin_db(uid):
        role_title = "SUPER ADMIN" if is_super(uid) else "ADMIN"
        text = f"👑 <b>{role_title} PANEL</b>\n\n━━━━━━━━━━━━━━━━━━\n📚 <b>Test bazasi</b>\n📝 Savollar: <b>{len(ALL_QUESTIONS)}</b>\n👥 Users: <b>{len(get_users())}</b>\n🛡 Adminlar: <b>{len(get_admins())}</b>\n━━━━━━━━━━━━━━━━━━\n🚀 Tizim tayyor"
    else:
        text = f"🎓 <b>QUIZ BOT</b>\n\n━━━━━━━━━━━━━━━━━━\n📚 Savollar: <b>{len(ALL_QUESTIONS)}</b>\n🏆 Reyting\n📊 Natijalar\n🌍 Umumiy test\n━━━━━━━━━━━━━━━━━━\n✅ Assalomu alaykum! Testni boshlash tugmasini bosing yoki qidiruv tizimidan foydalaning."

    await message.answer(text, parse_mode="HTML", reply_markup=main_menu(uid))


# ==================================================
# MUKAMMAL QIDIRUV TIZIMI (SEARCH ENGINE - 7 TA FAN)
# ==================================================
@dp.message(F.text == "🔍 Qidiruv")
async def start_search(message: Message, state: FSMContext):
    uid = message.from_user.id
    if uid != SUPER_ADMIN and not is_admin_db(uid) and not is_user(uid):
        await message.answer("⛔ Sizga ruxsat yo‘q")
        return

    await state.set_state(SearchState.wait_keyword)

    cancel_kb = ReplyKeyboardMarkup(
        keyboard=[[KeyboardButton(text="❌ Qidiruvni bekor qilish")]],
        resize_keyboard=True
    )

    await message.answer(
        "🔍 <b>Global Qidiruv Bo'limi</b>\n\n"
        "Barcha fanlar bazasidan qidirmoqchi bo'lgan kalit so'z yoki biron-bir jumlani yuboring.\n"
        "Bot sizga savol va uning to'g'ri javobini topib beradi.",
        parse_mode="HTML",
        reply_markup=cancel_kb
    )


@dp.message(SearchState.wait_keyword)
async def process_search(message: Message, state: FSMContext):
    uid = message.from_user.id
    keyword = message.text.strip()

    if keyword == "❌ Qidiruvni bekor qilish":
        await state.clear()
        return await message.answer("🏠 Qidiruv bekor qilindi.", reply_markup=main_menu(uid))

    if len(keyword) < 3:
        return await message.answer("⚠️ Qidiruv aniq bo'lishi uchun kamida 3 ta harf yoki so'z kiriting!")

    found_questions = []
    keyword_lower = keyword.lower()

    # Case-insensitive qisman qidiruv tizimi
    for q in ALL_QUESTIONS:
        if keyword_lower in q["question"].lower() or keyword_lower in q["subject"].lower():
            found_questions.append(q)

    if not found_questions:
        return await message.answer(
            f"❌ Afsuski, <b>\"{keyword}\"</b> bo'yicha hech qanday test topilmadi.\n\n"
            f"Boshqa so'z yozib ko'ring yoki bekor qiling.", parse_mode="HTML"
        )

    limit = 5  # Ekran to'lib ketmasligi va chiroyli chiqishi uchun limit
    text = f"🔍 <b>\"{keyword}\" bo'yicha topilgan eng yaqin natijalar ({len(found_questions)} ta):</b>\n\n"

    for i, q in enumerate(found_questions[:limit], start=1):
        text += (
            f"{i}. 📘 <b>Fan:</b> {q['subject']}\n"
            f"❓ <b>Savol:</b> {q['question']}\n"
            f"✅ <b>To'g'ri javob:</b> <u>{q['correct']}</u>\n"
            f"━━━━━━━━━━━━━━━━━━\n\n"
        )

    if len(found_questions) > limit:
        text += f"⚠️ <i>Yana {len(found_questions) - limit} ta mos keluvchi savol bor. Qidiruvni aniqroq so'z bilan qaytadan ko'ring.</i>"

    await message.answer(text, parse_mode="HTML")
    await state.clear()
    await message.answer("🏠 Asosiy menyuga qaytdingiz.", reply_markup=main_menu(uid))


# ==================================================
# ADMIN PANEL HANDLERLARI
# ==================================================
@dp.message(F.text == "👑 Admin Panel")
async def open_admin(message: Message):
    if not is_admin(message.from_user.id):
        return await message.answer("❌ Ruxsat yo‘q")
    await message.answer("👑 PANEL", reply_markup=admin_menu())


@dp.message(F.text == "👥 Users")
async def users_open(message: Message):
    if is_admin(message.from_user.id):
        await message.answer("👥 USERS", reply_markup=users_menu())


@dp.message(F.text == "🛡 Admin")
async def admins_open(message: Message):
    if not is_super(message.from_user.id):
        return await message.answer("❌ Faqat super admin ruxsati")
    await message.answer("🛡 ADMINS", reply_markup=admins_menu())


@dp.message(F.text == "📋 User List")
async def users_list(message: Message):
    if not is_admin(message.from_user.id): return
    users = get_users()
    if not users:
        return await message.answer("Bo‘sh")

    text = "👥 USERS LIST\n\n"
    for uid, fullname in users:
        text += f"👤 {fullname}\n🆔 {uid}\n━━━━━━━━━━\n"
    await message.answer(text)


@dp.message(F.text == "📋 Admin List")
async def admin_list(message: Message):
    if not is_admin(message.from_user.id): return
    text = f"🛡 ADMINS\n\n👑 SUPER\n{SUPER_ADMIN}\n\n"
    for uid in get_admins():
        text += f"🆔 {uid}\n"
        try:
            chat = await bot.get_chat(uid)
            text += f"👤 {chat.full_name}\n\n"
        except:
            text += "👤 Noma'lum\n\n"
    await message.answer(text)


# ==================================================
# ADMIN va USER QO'SHISH/O'CHIRISH (FSM)
# ==================================================
@dp.message(F.text == "➕ Admin")
async def add_admin_open(message: Message, state: FSMContext):
    if is_super(message.from_user.id):
        await state.set_state(AddAdmin.wait)
        await message.answer("Admin ID yuboring:")


@dp.message(AddAdmin.wait)
async def save_admin(message: Message, state: FSMContext):
    if message.text == "⬅ Back":
        await state.clear()
        return await message.answer("👑 ADMIN PANEL", reply_markup=admin_menu())
    if not message.text.isdigit():
        return await message.answer("Faqat raqamlardan iborat ID yuboring")

    uid = int(message.text)
    try:
        chat = await bot.get_chat(uid)
    except:
        return await message.answer("User hali botga kirmagan!\n1. Botga kirsin\n2. /start bossin")

    if uid == SUPER_ADMIN:
        await state.clear()
        return await message.answer("⚠ Bu SUPER ADMIN", reply_markup=admin_menu())

    add_admin_db(uid)
    save_result(uid, chat.full_name, 0, 0)
    await state.clear()
    await message.answer(f"✅ Admin qo‘shildi\n\n🆔 ID: {uid}\n👤 {chat.full_name}", reply_markup=admin_menu())


@dp.message(F.text == "❌ Del Admin")
async def del_admin_start(message: Message, state: FSMContext):
    if is_super(message.from_user.id):
        await state.set_state(DelAdmin.wait)
        await message.answer("O‘chiriladigan Admin ID yuboring:")


@dp.message(DelAdmin.wait)
async def save_del_admin(message: Message, state: FSMContext):
    if message.text == "⬅ Back":
        await state.clear()
        return await message.answer("🛡 ADMINS", reply_markup=admins_menu())
    if not message.text.isdigit():
        return await message.answer("Faqat ID yuboring")

    uid = int(message.text)
    if uid == SUPER_ADMIN:
        return await message.answer("❌ SUPER ADMIN o‘chirilmaydi")

    del_admin_db(uid)
    await state.clear()
    await message.answer(f"🗑 Admin o‘chirildi\n🆔 {uid}", reply_markup=admins_menu())


@dp.message(F.text == "➕ User")
async def add_user_open(message: Message, state: FSMContext):
    if is_admin(message.from_user.id):
        await state.set_state(AddUser.wait)
        await message.answer("User ID yuboring:")


@dp.message(AddUser.wait)
async def save_user(message: Message, state: FSMContext):
    if message.text == "⬅ Back":
        await state.clear()
        return await message.answer("👥 USERS", reply_markup=users_menu())
    if not message.text.isdigit():
        return await message.answer("ID kiriting")

    uid = int(message.text)
    if uid == SUPER_ADMIN:
        await state.clear()
        return await message.answer("⚠ Bu SUPER ADMIN", reply_markup=users_menu())

    try:
        chat = await bot.get_chat(uid)
    except:
        return await message.answer("User botga kirmagan! Avval botga kirib /start bossin.")

    if is_user(uid):
        await message.answer("⚠ User oldin qo‘shilgan")
    else:
        add_user_db(uid, chat.full_name)
        await message.answer(f"✅ User qo‘shildi\n\n🆔 ID: {uid}\n👤 {chat.full_name}")

    await state.clear()
    await message.answer("👥 USERS", reply_markup=users_menu())


@dp.message(F.text == "❌ Del User")
async def del_user_start(message: Message, state: FSMContext):
    if is_admin(message.from_user.id):
        await state.set_state(DelUser.wait)
        await message.answer("O‘chiriladigan ID yuboring:")


@dp.message(DelUser.wait)
async def save_del_user(message: Message, state: FSMContext):
    if message.text == "⬅ Back":
        await state.clear()
        return await message.answer("👥 USERS", reply_markup=users_menu())
    if not message.text.isdigit():
        return await message.answer("ID kiriting")

    uid = int(message.text)
    if uid == SUPER_ADMIN:
        return await message.answer("❌ SUPER ADMIN o‘chirib bo‘lmaydi")

    if is_user(uid):
        del_user_db(uid)
        await message.answer(f"🗑 User o‘chirildi\n🆔 {uid}")
    else:
        await message.answer("❌ User topilmadi")

    await state.clear()
    await message.answer("👥 USERS", reply_markup=users_menu())


# ==================================================
# NAVIGATSIYA TUGMALARI (BACK & HOME)
# ==================================================
@dp.message(F.text == "⬅ Back")
async def back(message: Message):
    uid = message.from_user.id
    if is_admin(uid):
        await message.answer("👑 PANEL", reply_markup=admin_menu())
    else:
        await message.answer("🏠 HOME", reply_markup=main_menu(uid))


@dp.message(F.text == "🏠 Home")
async def home(message: Message):
    await message.answer("HOME", reply_markup=main_menu(message.from_user.id))


# ==================================================
# QUIZ - TEST JARAYONI LOGIKASI
# ==================================================
@dp.message(F.text == "🚀 Testni boshlash")
async def test_menu(message: Message):
    uid = message.from_user.id
    if uid != SUPER_ADMIN and not is_admin_db(uid) and not is_user(uid):
        return await message.answer("⛔ Sizga ruxsat berilmagan.")

    counts = {}
    for q in ALL_QUESTIONS:
        counts[q["subject"]] = counts.get(q["subject"], 0) + 1

    text = "📚 Qaysi fandan test ishlaysiz?\n\n"
    kb = ReplyKeyboardBuilder()
    total = 0

    for subject, count in counts.items():
        total += count
        text += f"📘 {subject} — {count} ta savol\n"
        kb.row(KeyboardButton(text=f"📘 {subject}"))

    text += f"\n🌍 Barcha fanlar — {total} ta savol"
    kb.row(KeyboardButton(text="🌍 Barcha fanlar"))
    kb.row(KeyboardButton(text="🛑 Testni tugatish"))

    await message.answer(text, reply_markup=kb.as_markup(resize_keyboard=True))


@dp.message(F.text.startswith("📘 "))
async def subject_test(message: Message):
    user_id = message.from_user.id
    subject_name = message.text.replace("📘 ", "").split(" — ")[0]
    questions = [q for q in ALL_QUESTIONS if q["subject"] == subject_name]

    if not questions:
        return await message.answer("❌ Savollar topilmadi")

    random.shuffle(questions)
    questions = questions[:QUESTIONS_PER_TEST]

    users_test[user_id] = {
        "questions": questions, "index": 0, "correct": 0, "wrong": 0, "skipped": 0,
        "chat_id": message.chat.id, "start_time": time.time(), "answered": False,
        "subjects": {}, "full_name": message.from_user.full_name
    }
    await send_question(user_id)


@dp.message(F.text == "🌍 Barcha fanlar")
async def all_subjects_test(message: Message):
    user_id = message.from_user.id
    questions = ALL_QUESTIONS.copy()
    random.shuffle(questions)
    questions = questions[:QUESTIONS_PER_TEST]

    users_test[user_id] = {
        "questions": questions, "index": 0, "correct": 0, "wrong": 0, "skipped": 0,
        "chat_id": message.chat.id, "start_time": time.time(), "answered": False,
        "subjects": {}, "full_name": message.from_user.full_name
    }
    await send_question(user_id)


async def send_question(user_id: int):
    if user_id not in users_test: return
    data = users_test[user_id]

    if data["index"] >= len(data["questions"]):
        total = data["correct"] + data["wrong"] + data["skipped"]
        percent = round((data["correct"] / total) * 100, 1) if total > 0 else 0
        spent = int(time.time() - data["start_time"])

        subjects_text = "\n📚 Fanlar bo'yicha:\n\n"
        for sub, stat in data["subjects"].items():
            subjects_text += f"📘 {sub}\n✅ {stat['correct']} | ❌ {stat['wrong']} | ⏭ {stat['skipped']}\n\n"

        result_text = f"📊 TEST NATIJALARI\n\n👤 User ID: {user_id}\n📊 Jami savol: {total}\n✅ To'g'ri: {data['correct']}\n❌ Noto'g'ri: {data['wrong']}\n⏭ O'tkazilgan: {data['skipped']}\n🎯 Ball: {percent}%\n⏱ Vaqt: {spent // 60} daq {spent % 60} son\n{subjects_text}"

        user_results[user_id] = result_text
        save_result(user_id, data["full_name"], data["correct"], total)

        await bot.send_message(data["chat_id"], result_text, reply_markup=main_menu(user_id))
        users_test.pop(user_id, None)
        return

    q = data["questions"][data["index"]]
    answers = list(dict.fromkeys([clean_text(a)[:90] for a in q["answers"]]))

    if len(answers) < 2:
        data["index"] += 1
        return await send_question(user_id)

    correct_answer = clean_text(q["correct"])[:90]
    if correct_answer not in answers:
        answers[0] = correct_answer

    random.shuffle(answers)
    correct_index = answers.index(correct_answer)

    data["correct_index"] = correct_index
    data["answered"] = False

    text = f"📝 {data['index'] + 1}/{len(data['questions'])} | 🚀 {q['subject']}\n\n{q['question'][:250]}"
    await bot.send_poll(
        chat_id=data["chat_id"], question=text[:300], options=answers[:4],
        type=PollType.QUIZ, correct_option_id=correct_index, is_anonymous=False,
        explanation=f"✅ To'g'ri: {correct_answer}"
    )


@dp.poll_answer()
async def poll_answer_handler(poll_answer: PollAnswer):
    user_id = poll_answer.user.id
    if user_id not in users_test: return

    data = users_test[user_id]
    if data["answered"]: return
    data["answered"] = True

    q = data["questions"][data["index"]]
    subject = q["subject"]

    if subject not in data["subjects"]:
        data["subjects"][subject] = {"correct": 0, "wrong": 0, "skipped": 0}

    if not poll_answer.option_ids:
        data["skipped"] += 1
        data["subjects"][subject]["skipped"] += 1
    else:
        selected = poll_answer.option_ids[0]
        if selected == data["correct_index"]:
            data["correct"] += 1
            data["subjects"][subject]["correct"] += 1
        else:
            data["wrong"] += 1
            data["subjects"][subject]["wrong"] += 1

    data["index"] += 1
    await asyncio.sleep(2)
    await send_question(user_id)


# ==================================================
# STATISTIKA VA REYTINGLAR
# ==================================================
@dp.message(F.text == "📊 Natijalar")
async def results(message: Message):
    uid = message.from_user.id
    if uid != SUPER_ADMIN and not is_admin_db(uid) and not is_user(uid): return

    if uid not in user_results:
        return await message.answer("❌ Natija yo‘q")
    await message.answer(user_results[uid])


@dp.message(F.text == "🏆 Reyting")
async def leaderboard(message: Message):
    uid = message.from_user.id
    if uid != SUPER_ADMIN and not is_admin_db(uid) and not is_user(uid): return

    board = get_leaderboard()
    if not board:
        return await message.answer("❌ Reyting bo‘sh")

    text = "🏆 GLOBAL REYTING\n\n"
    for i, row in enumerate(board, start=1):
        name, correct, total = row[0], row[1], row[2]
        percent = round((correct / total) * 100, 1) if total else 0
        mark = " ⭐" if name == message.from_user.full_name else ""
        text += f"{i}. 👤 {name}{mark}\n✅ {correct}/{total}\n🎯 {percent}%\n\n"

    await message.answer(text)


@dp.message(F.text == "📚 Fanlar")
async def subjects(message: Message):
    counts = {}
    total = 0
    for q in ALL_QUESTIONS:
        counts[q["subject"]] = counts.get(q["subject"], 0) + 1
        total += 1

    text = "📚 MAVJUD FANLAR\n\n"
    for subject, count in counts.items():
        text += f"📘 {subject} — {count} ta savol\n"
    text += f"\n📝 Jami: {total} ta test"
    await message.answer(text)


@dp.message(F.text == "🛑 Testni tugatish")
async def stop_test(message: Message):
    user_id = message.from_user.id
    if user_id not in users_test:
        return await message.answer("❌ Siz hali test boshlamagansiz", reply_markup=main_menu(user_id))

    data = users_test[user_id]
    total = data["correct"] + data["wrong"] + data["skipped"]

    if total == 0:
        users_test.pop(user_id, None)
        return await message.answer("❌ Siz hech qanday savolga javob bermay testni tugatdingiz.",
                                    reply_markup=main_menu(user_id))

    percent = round((data["correct"] / total) * 100, 1)
    spent = int(time.time() - data["start_time"])

    subjects_text = "\n📚 Fanlar bo'yicha:\n\n"
    for sub, stat in data["subjects"].items():
        subjects_text += f"📘 {sub}\n✅ {stat['correct']} | ❌ {stat['wrong']} | ⏭ {stat['skipped']}\n\n"

    result_text = f"📊 TEST MUDDATIDAN OLDIN TUGATILDI\n\n👤 User ID: {user_id}\n📊 Jami bajarilgan savol: {total}\n✅ To'g'ri: {data['correct']}\n❌ Noto'g'ri: {data['wrong']}\n⏭ O'tkazilgan: {data['skipped']}\n🎯 Ball: {percent}%\n⏱ Sarflangan vaqt: {spent // 60} daq {spent % 60} son\n{subjects_text}"

    user_results[user_id] = result_text
    save_result(user_id, data["full_name"], data["correct"], total)

    users_test.pop(user_id, None)
    await message.answer(result_text, reply_markup=main_menu(user_id))


# ==================================================
# BOT ISHGA TUSHISH QISMI
# ==================================================
async def main():
    await dp.start_polling(bot)


if __name__ == "__main__":
    try:
        asyncio.run(main())
    except (KeyboardInterrupt, SystemExit):
        logging.info("Bot to'xtatildi!")